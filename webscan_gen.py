# coding: utf-8
# version: 2.0.1

from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE
from docx.shared import RGBColor, Pt, Cm, Inches
from docx.oxml.ns import nsdecls, qn
from pandas import read_excel
import pdfplumber as pdfp
import numpy as np
from datetime import datetime
import time
import traceback
import re
from webscan_doc import WebscanParag, WebscanTable
from _G import log_debug, log_error, log_warning, log_info
import _G
import util
import re

cur_table = WebscanTable()
cur_parag = cur_table.parag
start_scan_time=''
# Original plain text risk information that will be displayed
# if no translation available
original_risk_info = {}

def all_translated_risk_names(df):
  names = df.get(_G.XlsDevSheetName)[_G.XlsRiskColName].values
  names = [row for row in names if type(row) == str]
  return [name.split('\n') for name in names]

# 從 dev 中提取不同程度的風險個數
def scan_risk_cnt(dev_pdf):
  page = dev_pdf.pages[1]
  table = []
  sw, alert_cnt = False, 5

  # get risk cnt
  for pdf_table in page.extract_tables():
    for row in pdf_table:
      if _G.KwordRiskCnt in row:
        sw = True
      if sw and any(row):
        row = [x for x in row if x is not None]
        if alert_cnt:
          table.append([row[0], int(row[1])])
          alert_cnt -= 1

  # print(table, end='\n\n')
  log_debug("Risk Cnt Table:\n", table)
  return table

def get_original_risk_names(df, to_upper=False):
  tmp = df.get(_G.XlsDevSheetName)[_G.XlsRiskColName].values
  names = []
  for row in tmp:
    if type(row) != str:
      continue
    print(row, type(row))
    name = row.split('\n')[0]
    name_alias = name.split(' / ')
    _len = len(name_alias)
    last_name = None
    for i, name in enumerate(name_alias):
      if i == 0:
        last_name = name
        continue
      if abs(len(last_name) - len(name)) > len(last_name) * 0.5:
        name = last_name + name
      else:
        names.append(last_name.upper() if to_upper else last_name)
      last_name = name
    names.append(last_name.upper() if to_upper else last_name)
  return names

# 透過翻譯文件 (具所有的風險名稱)，依序對比出 dev 含有的風險，並取出成 list
# And collect original text in case of untranslated risk
def scan_risks(dev_pdf, df):
  risk_list = []

  names = get_original_risk_names(df, True)
  
  current_risk = None
  ori_risk_info = []
  cur_info_idx  = -1
  risk_info_idxs = [_G.KwordRiskLevel, _G.KwordRiskDesc, _G.KwordRiskImpact, _G.KwordRiskRecomd]
  last_seg = None

  for page in dev_pdf.pages[2:]:
    for pdf_table in page.extract_tables():
      # Extract valid risk name in row
      for row in pdf_table:
        for segments in row:
          if not segments:
            continue
          segments = re.split(r"[\r\n]+", segments)
          segments = [seg.strip() for seg in segments if seg.strip()]
          for segment in segments:
            log_debug(f"Current seg: {segment}")
            log_debug(f"Last seg: {last_seg}")
            risk_name = None
            if segment.isdigit(): # skip the page number
              continue
            if segment.upper() in names and segment not in risk_list:
              risk_name = segment
            elif segment == _G.KwordRisknameSuccessor and last_seg not in risk_list:
              risk_name = last_seg
            elif last_seg in risk_list and segment == _G.KwordRiskLevel:
              current_risk = last_seg
              ori_risk_info = ['', '', '', '']
              log_debug("Current risk info: ", current_risk)
            # Save original risk info
            if current_risk:
              if last_seg == _G.KwordRiskLevel:
                ori_risk_info[0] = segment
                cur_info_idx = -1
              elif segment in risk_info_idxs:
                cur_info_idx = risk_info_idxs.index(segment)
              elif segment in _G.KwordRiskInfoEnd:
                original_risk_info[current_risk] = ori_risk_info
                log_info("Ori risk pushed: ", current_risk, ori_risk_info)
                cur_info_idx = -1
                current_risk = None
              elif cur_info_idx < 0:
                pass
              else:
                ori_risk_info[cur_info_idx] += segment + ' '

            if risk_name:
              risk_list.append(risk_name)
            last_seg = segment

  log_debug("Risk list:", len(risk_list), risk_list, sep='\n',end='\n\n')
  return risk_list


def extract_valid_table_row(row):
  row = [word for word in row if word]
  if row[0].isdigit(): # page number
    return None
  return row

# 依序將 dev 的風險影響的 item 取出成 list
def scan_risk_range(dev_pdf, risk_cnts, risk_list):
  total = int(risk_cnts[0][1])
  cnt   = 0
  affect_range = []
  risk_ranges  = []
  have_next    = False    # flag current risk still have next affected item
  flag_cnt_done = False   # flag counting risk done
  cur_riskname = ''
  risk_idx = -1
  kword_depth = 0
  dict_risk_idx = {} # risk name and its index in risk ranges
  log_info("Risk List: ", risk_list)
  
  for page in dev_pdf.pages[2:]:
    for pdf_table in page.extract_tables():
      rows = []
      for cell in pdf_table:
        tmp = []
        for r in cell:
          tmp.extend([x.strip() for x in (r or '').split('\n') if x])
        if tmp:
          rows.append(tmp)
      for row in rows:
        if not row: # empty row
          continue
        log_debug("Row info:", row, have_next, flag_cnt_done)
        if not have_next and row[0] in risk_list:
          cur_riskname = row[0]
          log_debug("Risk section detected:", cur_riskname)
        if row and (row[0] in _G.KwordRiskRangeStart or row[-1] in _G.KwordRiskRangeStart) \
          or _G.KwordAffectedItems[0] in row:
          have_next = True
          continue
        if not have_next and not flag_cnt_done:
          continue
        
        # grab direct risk range
        if not flag_cnt_done:
          flag_range_end = (row[0] in _G.KwordRiskRangeEnd or row[-1] in _G.KwordRiskRangeEnd)
          log_debug("Flag count range end:", flag_range_end)
          if len(row) == 1 and row[0].isdigit(): # page number
            continue
          elif row[0] in risk_list or row[-1] in risk_list or flag_range_end:
            log_debug("Risk range pushed", cnt, affect_range)
            risk_ranges.append([cnt, affect_range])
            affect_range = []
            cnt = 0
            have_next = False
            cur_riskname = row[0]
            if len(cur_riskname) < 5:
              cur_riskname = row[-1]
            if flag_range_end:
              flag_cnt_done = True
              cur_riskname = None
              log_debug("Risk count done")
              print(dict_risk_idx)
          elif not flag_cnt_done and row[0] not in risk_list and len(row) > 1:
            if not row[1].isdigit():
              continue # not a valid affectd range entry, probably detected page header/footer
            log_debug(row)
            log_debug("Affact range pushed:", row[0])

            # expand items for webserver
            if re.search(_G.RegexWebServer, row[0], flags=re.IGNORECASE):
              affect_range.append(cur_riskname)
              dict_risk_idx[cur_riskname] = len(risk_ranges)
            else:
              affect_range.append(row[0])
            cnt = cnt + int(row[1])
            total = total - int(row[1])
        # expand affected webserver to paths if given
        else:
          if re.search(_G.RegexAffectExpandStop, row[0], flags=re.IGNORECASE):
            risk_ranges[risk_idx][1] = affect_range if affect_range else [_G.KwordWebServer]
            kword_depth = 0
            risk_idx = -1
            have_next = False
            continue
          if row[0] in dict_risk_idx:
            risk_idx = dict_risk_idx[row[0]]
            kword_depth = 1
            affect_range = []
          elif row[-1] in dict_risk_idx:
            risk_idx = dict_risk_idx[row[-1]]
            kword_depth = 1
            affect_range = []
          if risk_idx < 0:
            continue
          if kword_depth >= len(_G.KwordAffectedItems):
            kword_depth = 0
          elif row[0] in _G.KwordAffectedItems[kword_depth]:
            kword_depth += 1
            continue
          lines = []
          for segment in row:
            lines += [word.strip() for word in re.split(r"[\r\n]+", segment)]
          lines = [line for line in lines if line]
          log_debug("Current risk section:", risk_idx, risk_ranges[risk_idx], cur_riskname)
          log_debug(lines)
          for line in lines:
            if re.search(_G.RegexURI, line) and _G.Hostname in line:
              url = line.split(_G.Hostname)[-1].split()[0]
              affect_range.append(url)

  log_debug("Risk ranges:")
  log_debug(total,  end='\n\n') # sanity check, should be 0
  log_debug(risk_ranges, end='\n\n')
  print(risk_ranges)
  return risk_ranges

# 依照風險順序合併風險等級、影響 item 個數、影響 item
def merge_risk_detail(risk_cnt, risk_ranges):
  num = deepcopy(risk_cnt)
  details = deepcopy(risk_ranges)
  
  for i in range(len(risk_ranges)):
    for j in range(1, 5):
      if num[j][1]:
        num[j][1] = num[j][1] - details[i][0]
        details[i].insert(0, num[j][0])
        break
  
  log_debug("Risk details:")
  log_debug(details, end='\n\n')
  return details

# 未使用
def risk_describe(dev_pdf):
  sw = False
  total = []
  
  for page in dev_pdf.pages[2:]:
    text = page.extract_text() 
    if 'Alerts details' in text:
      sw = True
    if sw:
      tmp = text.split("\n")
      for row in tmp:
        total.append(row)


  risk_describes, risk_impacts, risk_recommendations = [], [], []

  for i, row in enumerate(total):
    if 'Description' in row:
      idx, end, tmp_str = 1, False, ""
      while not end:                                
        tmp_str = tmp_str + total[i+idx]
        
        if tmp_str[-1] != " ":
          tmp_str = tmp_str + " "
        
        idx = idx + 1            
        if 'Impact' in total[i+idx]:
          end = True

      risk_describes.append(tmp_str)

    if 'Impact' in row:
      idx, end, tmp_str = 1, False, ""
      while not end:            
        tmp_str = tmp_str + total[i+idx]    
        
        if tmp_str[-1] != " ":
          tmp_str = tmp_str + " "
        
        idx = idx + 1            
        if 'Recommendation' in total[i+idx]:
          end = True

      risk_impacts.append(tmp_str)    

    if 'Recommendation' in row:
      idx, end, tmp_str = 1, False, ""
      while not end:            
        tmp_str = tmp_str + total[i+idx]    
        
        if tmp_str[-1] != " ":
          tmp_str = tmp_str + " "
        
        idx = idx + 1            
        if 'Affected items' in total[i+idx]:
          end = True

      risk_recommendations.append(tmp_str)


  return risk_describes, risk_impacts, risk_recommendations

# 取出 owasp 十個類型的風險個數
def owasp_cnt(owa_pdf):
  owa_cnt = []

  page = owa_pdf.pages[1]
  fst_page = page.extract_text().split("\n")
  
  for i, row in enumerate(fst_page):
    if row[0] == '-': # version 1
      if 'No' in fst_page[i+1]:
        owa_cnt.append(0)
      else:
        num = fst_page[i+1].split(" ")
        owa_cnt.append(int(num[-1]))
      _G.TargetVersion = 0
    else: # version 2
      words = [w for w in row.split() if w]
      if len(words) > 2 and re.match(r"^A\d\d?$", words[1]):
        owa_cnt.append(int(words[0]))
        _G.TargetVersion = 1

  log_info("Detected PDF version:", _G.TargetVersion)
  log_debug("OWA count:")
  log_debug(owa_cnt, end='\n\n')
  return owa_cnt

# 透過 owasp 的風險個數和 dev 的風險清單，找出十個類型的風險名，並存成 list
def owasp_list(owa_pdf, owa_cnt, risk_list):
  owa_list = []
  if _G.TargetVersion == 0: # ancient code by unknown
    tmp, idx = [], 0
    verified_risk_list = [risk+' (verified)' for risk in risk_list]
    for page in owa_pdf.pages[1:]:
      for table in page.extract_tables():
        for row in table:
          if idx < 10:
            row = [x for x in row if x]
            if len(tmp) == owa_cnt[idx]:
              owa_list.append(tmp)
              tmp = []
              idx += 1
            # if row:
            #   print(row, row[0] in risk_list, row[0] in verified_risk_list)
            if row and (row[0] in risk_list or row[0] in verified_risk_list):
              row[0] = row[0].replace(' (verified)', '')
              tmp.append(row[0])
              # print(tmp)
          else:
            break
        if idx > 9:
          break
      if idx > 9:
        break
    log_debug("tmp list: ",tmp)
      # for tables
    # for pages
  elif _G.TargetVersion == 1:
    line_stack = []
    word_stack = []
    wsize_list = set()
    last_word  = ''
    cur_category = -1 # A1~A10
    
    # Detect entry and items by word size
    file = open('.log', 'w', encoding='utf8')
    for page in owa_pdf.pages[2:]:
      for word in page.extract_words(extra_attrs=['size']):
        wsize_list.add(float(word['size']))
        file.write(f"{word['size']} {word}\n")
    wsize_list = sorted(list(wsize_list))
    file.close()
    log_debug("Word sizes:", wsize_list)

    th_cate = np.average(wsize_list) + np.std(wsize_list)*2
    th_item = np.average(wsize_list) + np.std(wsize_list)*0.4-0.1
    log_info("Word size threshold:", th_cate, th_item)
    for page in owa_pdf.pages[2:]:
      for word in page.extract_words(extra_attrs=['size']):
        if word['size'] >= th_cate and re.match(r"^A\d\d?$", word['text'].strip()):
          owa_list.append([])
          cur_category += 1
          continue
        elif word['size'] >= th_item:
          word_stack.append(word['text'])
        elif word_stack:
          rname = ' '.join(word_stack)
          line_stack.append(rname)
          word_stack = []
          if rname in risk_list:
            owa_list[cur_category].append(rname)
        if last_word.lower().startswith('http') and word['text'].strip().lower() == 'verified':
          try:
            idx = line_stack.index('Impact')
          except ValueError:
            continue
          owa_list[cur_category].append(line_stack[idx-1])
          line_stack = []
        last_word = word['text'].strip()
  log_debug("line_stack: ",line_stack)
  
  log_debug("WORD list: ", word_stack)
  log_debug("OWASP list: ", owa_list)
  return owa_list

# 找出主要檢測的主要 host
def main_host(dev_pdf):
  for page in dev_pdf.pages:
    for table in page.extract_tables():
      for row in table:
        last_word = ''
        for word in row:
          if not word:
            continue
          word = f"{word}".strip()
          if not word:
            continue
          if last_word.upper() == _G.KwordMainhost.upper():
            return word
          last_word = word

# 具有 doc 前綴的 function，都是將資料寫入 docx 中
def doc_measurement_range(doc, df, url, dev_pdf):
  global  start_scan_time
  # start_scan_time = ''
  flag_time_found = False
  # Find start time of the scan
  for page in dev_pdf.pages:
    if flag_time_found:
      break
    for table in page.extract_tables():
      if flag_time_found:
        break
      for row in table:
        row = [word for word in row if word]
        try:
          if row[0] == _G.KwordScanStartTime and re.match(_G.RegexStartTime, row[1]):
            matches = re.match(_G.RegexStartTime, row[1]).groups()
            start_scan_time = f"{matches[0]}/{matches[1]}/{matches[2]}"
            print('start_scan_time : '+ start_scan_time)
            flag_time_found = True
            break
        except IndexError:
          pass

  table = doc.add_table(2, 3, style=_G.DocTableStyle)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  for i in range(3):
    table.cell(0, i).text = _G.DocMeasureRangeTitles[i]  
  table.cell(1,0).text = url
  table.cell(1,1).text = start_scan_time

  cur_table.set_content_font(table)
  cur_table.col_widths(table, 9, 5, 3)
  
  for p in doc.paragraphs:
    if 'measurementUrl' in p.text:
      p.text = p.text.replace('measurementUrl', url, 1)
      cur_parag.set_font(p, align=0)
    
    if 'docMeasurementRange' in p.text:
      cur_table.move_table_after(table, p)
      cur_parag.delete(p)
      log_info(f"Scan info table inserted: {table.cell(1,1).text}")
   
    
  
  if 'measurementUrl' in doc.tables[0].cell(0, 0).text:
    _font, _align = cur_parag.get_format_style(doc.tables[0].cell(0, 0).paragraphs)
    fsize  = _font.size if _font  else Pt(16)
    fbold  = _font.bold if _font  else True
    falign = _align     if _align else 2
    doc.tables[0].cell(0, 0).text = doc.tables[0].cell(0, 0).text.replace('measurementUrl', url)
    cur_parag.set_font(doc.tables[0].cell(0, 0).paragraphs, size=fsize, bold=fbold, align=falign)


def doc_risk_cnt(doc, risk_cnts):
  title = deepcopy(_G.DocRiskCountTitles)
  
  for i in range(1, 5):
    title[1].append(risk_cnts[i][1])

  table = doc.add_table(2, 5, style=_G.DocTableStyle)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  for i in range(2):
    for j in range(5):
      table.cell(i, j).text = str(title[i][j])

  cur_table.set_content_font(table)
  
  for p in doc.paragraphs:
    if _G.RepRiskCntList in p.text:
      cur_table.move_table_after(table, p)
      cur_parag.delete(p)
        
    row = p.text
    
    if any(_repw in row for _repw in _G.RepRiskLevel):
      row = row.replace(_G.RepRiskLevel[0], str(title[1][1]+title[1][2]+title[1][3]+title[1][4]), 1)
      row = row.replace(_G.RepRiskLevel[1], str(title[1][1]), 1)
      row = row.replace(_G.RepRiskLevel[2], str(title[1][2]), 1)
      row = row.replace(_G.RepRiskLevel[3], str(title[1][3]), 1)
      row = row.replace(_G.RepRiskLevel[4], str(title[1][4]), 1)
      p.text = row
      cur_parag.set_font(p, align=0)


def doc_risk_list(doc, df, risk_list, risk_details):

  names = all_translated_risk_names(df)
  table = doc.add_table(1+len(risk_list), 4, style=_G.DocTableStyle)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  for i in range(1+len(risk_list)):
    log_debug(f"Writing risk list #{i}")
    for j in range(4):
      cell = table.cell(i, j)
      if not i:
        cell.text = _G.DocRiskListTitles[j]
      else:
        if not j:
          cell.text = str(i)
        elif j == 1:
          cell.text = _G.DocRiskLevelTrans[risk_details[i-1][0]]
        elif j == 2:
          name = risk_list[i-1]
          cell.text = risk_name_translate(name, names)
        elif j == 3:
          cell.text = str(risk_details[i-1][1])

    cur_table.set_content_font(table)
    for i in range(len(risk_list)):
      cell = table.cell(i+1, 2)
      cur_parag.set_font(cell.paragraphs, align=0)
    
    cur_table.col_widths(table, 2, 3, 10, 2)
    
    for parag in doc.paragraphs:
      if _G.RepRiskList in parag.text:
        cur_table.move_table_after(table, parag)
        cur_parag.delete(parag)
        break


def doc_owasp_list(doc, df, owa_list):
  data = df.get(_G.XlsOwaspSheetName).values
  names = all_translated_risk_names(df)
  table = doc.add_table(11, 3, style=_G.DocTableStyle)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  for i in range(3):
    for j in range(11):
      cell = table.cell(j, i)

      if not j:
        cell.text = _G.DocOwaspListTitle[i]
      else: 
        if not i:
          cell.text = data[j-1][0]

        elif i == 1:
          cell.text = data[j-1][1].split('\n')[1]

        elif i == 2:
          if len(owa_list) > j and owa_list[j-1]:
            name = owa_list[j-1][0]
            tmp = risk_name_translate(name, names)
            for z in owa_list[j-1][1:]:
              name = risk_name_translate(z, names)
              if name not in tmp:
                tmp = tmp + '\n' + name
            cell.text = tmp
          else:
            cell.text = '-'    
                  
  cur_table.set_content_font(table, align=0)
  for i in range(11):
    cell = table.cell(i, 0)
    cur_parag.set_font(cell.paragraphs, align=1, bold=True if not i else False)
      
  cur_table.col_widths(table, 1.5, 6, 9.5)
  
  for parag in doc.paragraphs:
    if _G.RepOwaspList in parag.text:
      cur_table.move_table_after(table, parag)
      cur_parag.delete(parag)
      break


def doc_owasp_risk(doc, df, owa_risks):

  names = all_translated_risk_names(df) 
  des = []  
  for item in owa_risks:
    des.append(risk_name_translate(item, names))

  for p in doc.paragraphs:
    if _G.RepOwaspRisk in p.text:
      p.text = p.text.replace(_G.RepOwaspRisk, '、'.join(des))
      cur_parag.set_font(p, align=0)
      break


def doc_risk_describe(doc, df, risk_list, risk_details):
  data = df.get(_G.XlsDevSheetName).values
  names = all_translated_risk_names(df) 

  for location in doc.paragraphs:
    if _G.RepRiskDesc in location.text:            
      for i, risk in enumerate(reversed(risk_list)):
          idx, front = risk_name_translate(risk, names, idx=True)

          title = cur_parag.insert_paragraph_after(location, front)
          cur_parag.set_font(title, size=Pt(_G.Config['RiskTitleFontSize']), bold=True, align=0)

          for j in reversed(range(5)):
            parag = cur_parag.insert_paragraph_after(title, _G.DocRiskDetails[j], style=_G.DocParagStyle)
            
            cur_parag.create_list(parag, _G.DocListTypeUnordered)
            cur_parag.set_font(parag, size=Pt(_G.Config['NormalFontSize']), bold=True, align=0)

            if not j:
              if idx < 0:
                level = original_risk_info[front][0]
              else:
                level = '資訊風險' if '資' in data[idx][1] else data[idx][1]
              run = parag.add_run(level)

            # 風險內容概述
            elif j == 1:
              if idx < 0:
                desc = original_risk_info[front][1]
              else:
                desc = data[idx][3] if isinstance(data[idx][3], str) else '-'
              run = parag.add_run('\n' + desc)

            # 衝擊
            elif j == 2:
              if idx < 0:
                impact = original_risk_info[front][2]
              else:
                tmp = data[idx][4].split('\n', 1) if isinstance(data[idx][4], str) else '-'
                if is_all_ascii(tmp[0]) and len(tmp) > 1:
                  del tmp[0]  
                impact = '\n'.join(tmp)
              if not impact.strip():
                impact = '-'
              run = parag.add_run('\n' + impact)
            # 影響範圍
            elif j == 3:
              flag_added = False
              for z in risk_details[len(risk_list)-i-1][2]:
                run = parag.add_run('\n' + z)
                run.font.size = Pt(_G.Config['NormalFontSize'])
                run.font.name = _G.Config['FontEnglish']
                flag_added = True
              if not flag_added:
                run = parag.add_run('\n-')
                run.font.size = Pt(_G.Config['NormalFontSize'])
                run.font.name = _G.Config['FontEnglish']
            # 建議
            elif j == 4:
              if idx < 0:
                recomman = original_risk_info[front][3]
              else:
                tmp = data[idx][5].split('\n', 1) if isinstance(data[idx][5], str) else '-'
                if is_all_ascii(tmp[0]) and len(tmp) > 1:
                  del tmp[0]
                recomman = '\n'.join(tmp)
              run = parag.add_run('\n' + recomman)
              if i:
                run.add_break(WD_BREAK.PAGE)

            run.font.size = Pt(_G.Config['NormalFontSize'])
            run.font.name = _G.Config['FontEnglish']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), _G.Config['FontChinese'])
          
      cur_parag.delete(location)

# 將英文風險名，透過翻譯 excel 轉成中文
def risk_name_translate(name, names, idx=False):
  if idx:
    for i, ch in enumerate(names):
      if name.upper() in ch[0].upper() and len(ch) > 1:
        return i, ch[1].replace('\0','')
      elif name.upper() in ch[0].upper():
        return i, ch[0]
    return -1, name.replace('\0','')
  else:
    for ch in names:
      if name.upper() in ch[0].upper() and len(ch) > 1:
        return ch[1].replace('\0','')
      elif name.upper() in ch[0].upper():
        return ch[0].replace('\0','')
    # No translation found
    return name.replace('\0','')

# 判斷是否全為 ascii
def is_all_ascii(strs):
  for _char in strs:
    if not '\0' <= _char <= '~':
      return False
  return True

def fill_company_name(doc, name, abbr):
  flag_first_page = True
  for parag in doc.paragraphs:
    if _G.RepCompanyName in parag.text:
      parag.text = parag.text.replace(_G.RepCompanyName, name)
      if flag_first_page:
        cur_parag.set_font(parag, Pt(28), True, None)
        flag_first_page = False
      else:
        cur_parag.set_font(parag, align=None) # Force same font in the paragraph
    if _G.RepCompanyAbbr in parag.text:
      parag.text = parag.text.replace(_G.RepCompanyAbbr, abbr)
      cur_parag.set_font(parag, align=None) # Force same font in the paragraph

def fill_doc_date(doc):
  for table in doc.tables:
    width, height = len(table.columns), len(table.rows)
  for i in range(height):
    for j in range(width):
      try:
        if table.cell(i, j).text == _G.KwordDocDate and len(table.cell(i+1, j).text.strip()) < 2:
          table.cell(i+1, j).text = f"{datetime.now().year}/{datetime.now().month}/{datetime.now().day}"
          cur_table.set_content_font(table, None, False)
          return
      except IndexError:
        pass

def generate_report(dev_pdffile, owa_pdffile, doc_file, xls_file, save_path, **kwargs):
  global cur_table, cur_parag
  cur_table = WebscanTable()
  cur_parag = cur_table.parag

  dev_pdf = pdfp.open(dev_pdffile)
  owa_pdf = pdfp.open(owa_pdffile)
  
  start_t = time.time()

  _G.append_pipe_message(_G.MsgInit)
  doc = Document(doc_file)
  df = read_excel(xls_file, sheet_name=None)
  
  fill_company_name(doc, kwargs.get('company_name', _G.RepCompanyName), kwargs.get('company_abbr', _G.RepCompanyAbbr))
  fill_doc_date(doc)
  log_debug(f"Time for init: {time.time() - start_t}")
  
  _G.append_pipe_message(_G.MsgHost)
  start_t = time.time()
  url = main_host(dev_pdf)
  _G.Hostname = re.split(_G.RegexURI,url)[-1].split('/')[0]
  log_info("Hostname:", _G.Hostname)
  log_debug(f"Time for search host: {time.time() - start_t}")
  
  _G.append_pipe_message(_G.MsgRisks)
  start_t = time.time()
  risk_cnts = scan_risk_cnt(dev_pdf)
  log_debug(f"Time for count risk: {time.time() - start_t}")
  
  start_t = time.time()
  risk_list = scan_risks(dev_pdf, df)
  log_debug(f"Time for list risk: {time.time() - start_t}")
 
  _G.append_pipe_message(_G.MsgRiskRange)
  start_t = time.time()
  risk_ranges = scan_risk_range(dev_pdf, risk_cnts, risk_list)
  log_debug(f"Time for risk range: {time.time() - start_t}")
  
  start_t = time.time()
  risk_details = merge_risk_detail(risk_cnts, risk_ranges)
  log_debug(f"Time for risk details: {time.time() - start_t}")

  start_t = time.time()
  _G.append_pipe_message(_G.MsgOwaList)
  owa_cnt = owasp_cnt(owa_pdf)
  log_debug(f"Time for owa_cnt: {time.time() - start_t}")
  
  start_t = time.time()
  _G.append_pipe_message(_G.MsgOwaRisks)
  owa_list = owasp_list(owa_pdf, owa_cnt, risk_list)
  owa_risks = set(list(util.flatten(owa_list)))
  log_debug("OWASP risks:", owa_risks)
  log_debug(f"Time for owa_risks: {time.time() - start_t}")
  
  _G.append_pipe_message(_G.MsgDocGen)
  start_t = time.time()
  doc_measurement_range(doc, df, url, dev_pdf)
  log_debug(f"Time for doc_measurement_range: {time.time() - start_t}")
  
  start_t = time.time()
  doc_risk_cnt(doc, risk_cnts)
  log_debug(f"Time for doc_risk_cnt: {time.time() - start_t}")
  
  start_t = time.time()
  doc_risk_list(doc, df, risk_list, risk_details)
  log_debug(f"Time for doc_risk_list: {time.time() - start_t}")
  
  start_t = time.time()
  doc_owasp_list(doc, df, owa_list)
  log_debug(f"Time for doc_owasp_list: {time.time() - start_t}")
  
  start_t = time.time()
  doc_owasp_risk(doc, df, owa_risks)
  log_debug(f"Time for doc_owasp_risk: {time.time() - start_t}")
  
  start_t = time.time()
  doc_risk_describe(doc, df, risk_list, risk_details)
  log_debug(f"Time for doc_risk_describe: {time.time() - start_t}")

  start_t = time.time()
  version_information(doc, kwargs.get('company_name', _G.RepCompanyName), kwargs.get('date', _G.Repdate))  
  log_debug(f"version_information: {time.time() - start_t}")

  save_ok = False
  flag_waiting = False
  while not save_ok:
    if _G.PipeSubInfo and _G.PipeSubInfo[0] == _G.MsgPipeContinue:
      flag_waiting = False
      _G.PipeSubInfo.popleft()
    elif flag_waiting:
      time.sleep(0.1)
      continue
    try:
      doc.save(save_path)
      save_ok = True
    except PermissionError as err:
      log_warning(err, traceback.format_exc())
      flag_waiting = True
      _G.PipeWarning.append(_G.MsgPipeWarnTargetOpened)
  
  if _G.RepMeasureURL in doc.tables[0].cell(0, 0).text:
    doc.tables[0].cell(0, 0).text = doc.tables[0].cell(0, 0).text.replace(_G.RepMeasureURL, url)
      
  for p in doc.paragraphs:
    if _G.RepMeasureURL in p.text:
      p.text = p.text.replace(_G.RepMeasureURL, url, 1)
  
  log_info("Completed")

def silent_generate_report(dev_pdf, owa_pdf, doc_file, xls_file, save_path, **kwargs):
  try:
    generate_report(dev_pdf, owa_pdf, doc_file, xls_file, save_path, **kwargs)
  except Exception as err:
    err_info = traceback.format_exc()
    _G.PipeError.append([err, err_info])
    log_error(f"An error occurred during generating report!\n{err_info}")

def version_information(doc, name, date):
   
  table = doc.add_table(rows=6, cols=4, style='Table Grid')
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER    # 第一行表格水平居中
  table.cell(0,0).merge(table.cell(0,3))
  table.cell(0,0).text ="版本資訊"
  table.cell(1,0).text ="姓名"
  #
  table.cell(1,1).text = name+"_網站弱點掃描"
  table.cell(1,1).merge(table.cell(1,3))
  # # table.cell(1,2).text = '' 
  # # table.cell(1,3).text = '' 
  table.cell(2,0).text = '撰寫者' 
  table.cell(2,1).text='黃湘'
  table.cell(2,2).text = '撰寫日期'
  table.cell(2,3).text = date
  
  table.cell(3,0).text = '版本清單'
  table.cell(3,0).merge(table.cell(3,3))
  table.cell(4,0).text = '版本編號'
  table.cell(4,1).text = '版本日期'
  table.cell(4,2).text = '修改者'
  table.cell(4,3).text = '說明'
  table.cell(5,0).text = 'V1.0'
  table.cell(5,1).text = start_scan_time
  table.cell(5,2).text = '黃湘'
  table.cell(5,3).text = ''
  for p in doc.paragraphs:
    
    if 'version_info' in p.text:
      cur_table.move_table_after(table, p)
      cur_parag.delete(p)
      log_info(f"Scan info table inserted: {table.cell(1,1).text}")
  cur_table.set_content_font(table)
  cur_table.col_widths(table, 2.5, 4, 2.5)

if __name__ == "__main__":
  doc_file = "C:/Users/dtyty/桌面/SystexWebScanReportGen/templates/網站弱點掃描報告.docx"
  doc = Document(doc_file)
  # version_information(doc, kwargs={'company_name': 'aa', 'company_abbr': 'company_abbr', 'date':'aaaa'})
  # doc.save('aa.docx')
  # for row in doc.tables[2].rows:
  #   for cell in row.cells:
  #     print(cell.text)